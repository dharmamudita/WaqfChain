from docx import Document

doc = Document('proposal_bp_final_rapi_BAB.docx')

# Fix: "2.1 Strategi Bisnis" should be "2.2 Strategi Bisnis"
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("2.1 Strategi Bisnis"):
        for run in para.runs:
            run.text = run.text.replace("2.1 Strategi Bisnis", "2.2 Strategi Bisnis")
        print(f"Fixed paragraph [{i}]: {para.text.strip()}")
        break

# Fix: "2.2 Strategi Pemasaran" should be "2.3 Strategi Pemasaran"  
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("2.2 Strategi Pemasaran"):
        for run in para.runs:
            run.text = run.text.replace("2.2 Strategi Pemasaran", "2.3 Strategi Pemasaran")
        print(f"Fixed paragraph [{i}]: {para.text.strip()}")
        break

doc.save('proposal_bp_final_rapi_BAB.docx')
print("Fixes saved!")

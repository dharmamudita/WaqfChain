from docx import Document

doc = Document('proposal_bp_final_rapi_BAB.docx')

replacements = {
    "Latar Belakang": "1.1 Latar Belakang",
    "Visi dan Misi": "1.2 Visi dan Misi",
    "Deskripsi Bisnis yang Ditawarkan": "1.3 Deskripsi Bisnis yang Ditawarkan",
    "Analisis Pasar": "2.1 Analisis Pasar",
    "Strategi Bisnis": "2.2 Strategi Bisnis",
    "Strategi Pemasaran": "2.3 Strategi Pemasaran",
    "Biaya Operasional": "3.2 Biaya Operasional",
    "Analisis Titik Impas": "3.3 Analisis Titik Impas (BEP)",
    "Perhitungan Kelayakan Usaha dengan Pendekatan Ekonomi": "3.4 Perhitungan Kelayakan Usaha",
    "Permodalan (Modal Pribadi / Dari Lembaga)": "4.1 Permodalan (Modal Pribadi/Lembaga)",
    "Manajemen Bisnis": "4.2 Manajemen Bisnis",
    "Akad akad yang digunakan": "4.3 Akad-Akad yang Digunakan"
}

# We need to make sure we only replace in the content, not in the TOC.
# TOC ends around paragraph 60.

for i, para in enumerate(doc.paragraphs):
    if i < 60:
        continue
    
    text = para.text.strip()
    
    for old, new in replacements.items():
        if text == old:
            # Clear runs and insert new text to preserve styles
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new
            print(f"Replaced: '{old}' -> '{new}'")

# Handle BAB V / Kesimpulan specifically
for i, para in enumerate(doc.paragraphs):
    if i < 60:
        continue
    text = para.text.strip()
    if text == "PENUTUP":
        # Look ahead a bit to see if Kesimpulan exists
        has_kesimpulan = False
        for j in range(1, 5):
            if i + j < len(doc.paragraphs) and "Kesimpulan" in doc.paragraphs[i+j].text:
                has_kesimpulan = True
                break
        
        if not has_kesimpulan:
            # Insert "5.1 Kesimpulan" after "PENUTUP"
            from copy import deepcopy
            new_p = deepcopy(para._element)
            para._element.addnext(new_p)
            from docx.text.paragraph import Paragraph
            new_para = Paragraph(new_p, para._parent)
            for run in new_para.runs:
                run.text = ''
            if new_para.runs:
                new_para.runs[0].text = "5.1 Kesimpulan"
                new_para.runs[0].font.bold = True
            print("Inserted '5.1 Kesimpulan'")
        break

doc.save('proposal_bp_final_rapi_BAB_Lengkap.docx')
print("File updated with numbered subheadings and saved as proposal_bp_final_rapi_BAB_Lengkap.docx!")

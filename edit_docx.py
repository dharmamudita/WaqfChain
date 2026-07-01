from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy

doc = Document('proposal_bp_final_rapi.docx')

# ============================================================
# STEP 1: UBAH DAFTAR ISI (paragraphs 38-49)
# ============================================================
# Mapping dari teks lama ke teks baru
toc_replacements = {
    38: ("GAMBAR UNIT BISNIS", "BAB I PENDAHULUAN"),
    39: ("A. Latar Belakang", "   1.1 Latar Belakang"),
    40: ("B. Visi dan Misi", "   1.2 Visi dan Misi"),
    41: ("C. Deskripsi Bisnis yang Ditawarkan", "   1.3 Deskripsi Bisnis yang Ditawarkan"),
    42: ("D. Aspek Perencanaan", "BAB II ASPEK PERENCANAAN DAN PEMASARAN"),
    43: ("E. Strategi Bisnis", "   2.1 Strategi Bisnis"),
    44: ("F. Strategi Pemasaran", "   2.2 Strategi Pemasaran"),
    45: ("G. Aspek Finansial", "BAB III ASPEK FINANSIAL"),
    46: ("H. Aspek Kesesuaian Syariah", "BAB IV ASPEK KESESUAIAN SYARIAH"),
    47: ("PENUTUP", "BAB V PENUTUP"),
}

# Kita perlu menambahkan sub-bab baru untuk BAB II, III, IV
# Tapi karena menambah paragraf di tengah itu rumit, kita ubah teks yang ada saja
# dan tambahkan sub-items sebagai baris baru setelahnya

for idx, (old_prefix, new_text) in toc_replacements.items():
    para = doc.paragraphs[idx]
    # Pertahankan tab dan nomor halaman yang sudah ada
    full_text = para.text
    # Ambil page number (setelah tab)
    parts = full_text.split('\t')
    page_num = parts[-1].strip() if len(parts) > 1 else ''
    
    # Update setiap run dalam paragraf
    # Kita clear semua runs lalu set text di run pertama
    if para.runs:
        # Simpan formatting dari run pertama
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None
        
        # Set teks baru dengan tab dan page number
        for run in para.runs:
            run.text = ''
        
        # Tentukan apakah ini heading BAB (bold) atau sub-item
        is_bab = new_text.startswith("BAB")
        first_run.text = new_text + '\t' + page_num
        first_run.font.bold = is_bab
        if font_name:
            first_run.font.name = font_name
        if font_size:
            first_run.font.size = font_size

print("Daftar Isi updated!")

# ============================================================
# STEP 2: Tambahkan sub-bab yang belum ada di Daftar Isi
# Kita perlu insert paragraf baru setelah beberapa entry
# ============================================================
# Untuk BAB II, perlu insert "2.1 Analisis Pasar" sebelum "2.1 Strategi Bisnis"
# Kita insert paragraf baru setelah paragraf 42 (BAB II)

def insert_paragraph_after(paragraph, text, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    # Clear and set new text
    for run in new_para.runs:
        run.text = ''
    if new_para.runs:
        new_para.runs[0].text = text
    else:
        new_run = new_para.add_run(text)
    return new_para

# Insert "2.1 Analisis Pasar" after BAB II line
new_para = insert_paragraph_after(doc.paragraphs[42], "   2.1 Analisis Pasar\t5")
if new_para.runs:
    new_para.runs[0].font.bold = False

# After BAB III, we need sub items for 3.1-3.4
# But first let's add sub-items after BAB III (index 45, but shifted by 1 = 46 now)
bab3_subs = [
    "   3.1 Permodalan\t11",
    "   3.2 Biaya Operasional\t12",
    "   3.3 Analisis Titik Impas (BEP)\t13",
    "   3.4 Perhitungan Kelayakan Usaha\t14",
]

# Insert after BAB III paragraph
# After the first insert, indices shifted. We need to re-find BAB III
# Let's find it by text
bab3_para = None
for i, p in enumerate(doc.paragraphs):
    if "BAB III ASPEK FINANSIAL" in p.text:
        bab3_para = p
        break

if bab3_para:
    prev = bab3_para
    for sub_text in bab3_subs:
        prev = insert_paragraph_after(prev, sub_text)
        if prev.runs:
            prev.runs[0].font.bold = False

# Insert sub-items after BAB IV
bab4_subs = [
    "   4.1 Permodalan (Modal Pribadi/Lembaga)\t15",
    "   4.2 Manajemen Bisnis\t17",
    "   4.3 Akad-Akad yang Digunakan\t18",
]

bab4_para = None
for i, p in enumerate(doc.paragraphs):
    if "BAB IV ASPEK KESESUAIAN SYARIAH" in p.text:
        bab4_para = p
        break

if bab4_para:
    prev = bab4_para
    for sub_text in bab4_subs:
        prev = insert_paragraph_after(prev, sub_text)
        if prev.runs:
            prev.runs[0].font.bold = False

# Insert "5.1 Kesimpulan" after BAB V PENUTUP
bab5_para = None
for i, p in enumerate(doc.paragraphs):
    if "BAB V PENUTUP" in p.text:
        bab5_para = p
        break

if bab5_para:
    new_p = insert_paragraph_after(bab5_para, "   5.1 Kesimpulan\t20")
    if new_p.runs:
        new_p.runs[0].font.bold = False

print("Sub-items inserted!")

# ============================================================
# STEP 3: UBAH HEADING DI HALAMAN KONTEN
# ============================================================
# "GAMBAR UNIT BISNIS" -> "BAB I PENDAHULUAN"
# Dan tambahkan heading BAB II, III, IV, V di posisi yang tepat

content_heading_replacements = {
    "GAMBAR UNIT BISNIS": "BAB I\nPENDAHULUAN",
}

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Ganti "GAMBAR UNIT BISNIS" menjadi "BAB I PENDAHULUAN"
    if text == "GAMBAR UNIT BISNIS" and i > 50:  # Yang di konten, bukan di TOC
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = "BAB I"
        # Tambah paragraf "PENDAHULUAN" setelahnya
        new_p = insert_paragraph_after(para, "PENDAHULUAN")
        if new_p.runs:
            new_p.runs[0].font.bold = True
            new_p.runs[0].font.size = para.runs[0].font.size if para.runs[0].font.size else Pt(14)
        # Copy alignment
        if para.alignment is not None:
            new_p.alignment = para.alignment
        break

# Cari heading-heading konten yang perlu diubah jadi BAB
# "Aspek Perencanaan" di konten -> tambah "BAB II" sebelumnya
content_bab_markers = [
    ("Aspek Perencanaan", "BAB II\nASPEK PERENCANAAN DAN PEMASARAN"),
    ("Aspek Finansial", "BAB III\nASPEK FINANSIAL"),
    ("Aspek Kesesuain Syariah", "BAB IV\nASPEK KESESUAIAN SYARIAH"),
]

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    for old_heading, new_heading in content_bab_markers:
        if text == old_heading and i > 50:
            lines = new_heading.split('\n')
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = lines[0]
                para.runs[0].font.bold = True
            # Add subtitle
            if len(lines) > 1:
                new_p = insert_paragraph_after(para, lines[1])
                if new_p.runs:
                    new_p.runs[0].font.bold = True
                    if para.runs[0].font.size:
                        new_p.runs[0].font.size = para.runs[0].font.size
                if para.alignment is not None:
                    new_p.alignment = para.alignment
            break

# Find and update PENUTUP heading in content
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text == "PENUTUP" and i > 50:
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = "BAB V"
            para.runs[0].font.bold = True
        new_p = insert_paragraph_after(para, "PENUTUP")
        if new_p.runs:
            new_p.runs[0].font.bold = True
            if para.runs[0].font.size:
                new_p.runs[0].font.size = para.runs[0].font.size
        if para.alignment is not None:
            new_p.alignment = para.alignment
        break

print("Content headings updated!")

# ============================================================
# SAVE
# ============================================================
doc.save('proposal_bp_final_rapi_BAB.docx')
print("File saved successfully as proposal_bp_final_rapi_BAB.docx!")
